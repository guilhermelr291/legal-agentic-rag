"""DOCX text extractor using python-docx."""

from docx import Document

from src.extractors.base import ExtractionResult, TextExtractor


class DOCXExtractor(TextExtractor):
    """Extract text and metadata from DOCX files."""

    def extract(self, file_path: str) -> ExtractionResult:
        """Extract text and metadata from a DOCX file.

        Args:
            file_path: Path to the DOCX file

        Returns:
            ExtractionResult with text, paragraph metadata, and file info
        """
        try:
            doc = Document(file_path)

            # Extract paragraphs with their properties
            paragraphs: list[dict] = []
            full_text_parts: list[str] = []
            current_char = 0

            for para_idx, para in enumerate(doc.paragraphs):
                para_text = para.text.strip()
                if not para_text:
                    continue

                para_info: dict = {
                    "index": para_idx,
                    "text": para_text,
                    "start_char": current_char,
                    "end_char": current_char + len(para_text),
                }

                # Extract style information if available
                if para.style and para.style.name:
                    para_info["style"] = para.style.name

                # Check if paragraph looks like a heading
                if para.style and para.style.name:
                    style_lower = para.style.name.lower()
                    if "heading" in style_lower or "title" in style_lower:
                        para_info["is_heading"] = True
                        para_info["heading_level"] = self._extract_heading_level(para.style.name)

                # Check for numbering/bullet
                if para._p is not None:
                    # Check for numbering properties
                    num_pr = para._p.find(
                        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr"
                    )
                    if num_pr is not None:
                        para_info["has_numbering"] = True

                paragraphs.append(para_info)
                full_text_parts.append(para_text)
                current_char += len(para_text) + 1  # +1 for newline

            full_text = "\n".join(full_text_parts)

            # Build page-like structure (DOCX doesn't have explicit pages)
            # We approximate by grouping paragraphs into "pages" of ~3000 chars
            pages = self._build_page_boundaries(paragraphs, full_text)

            # Extract metadata
            metadata: dict = {
                "source": file_path,
                "paragraph_count": len(doc.paragraphs),
                "extracted_paragraphs": len(paragraphs),
            }

            # Document properties from core properties
            if doc.core_properties:
                core_props = doc.core_properties
                if core_props.author:
                    metadata["author"] = core_props.author
                if core_props.title:
                    metadata["title"] = core_props.title
                if core_props.subject:
                    metadata["subject"] = core_props.subject
                if core_props.created:
                    metadata["created"] = str(core_props.created)
                if core_props.modified:
                    metadata["modified"] = str(core_props.modified)

            return ExtractionResult(
                text=full_text,
                pages=pages,
                metadata=metadata,
                success=True,
            )

        except FileNotFoundError:
            return self._create_error_result(f"File not found: {file_path}")
        except PermissionError:
            return self._create_error_result(f"Permission denied: {file_path}")
        except Exception as e:
            return self._create_error_result(f"DOCX extraction failed: {e!s}")

    def _extract_heading_level(self, style_name: str) -> int:
        """Extract heading level from style name.

        Args:
            style_name: The style name (e.g., "Heading 1", "Heading 2")

        Returns:
            Heading level (1-9), defaults to 1 if not determinable
        """
        style_lower = style_name.lower()
        if "heading" in style_lower:
            # Try to extract number from "Heading X"
            parts = style_name.split()
            for part in parts:
                if part.isdigit():
                    return int(part)
        return 1

    def _build_page_boundaries(
        self, paragraphs: list[dict], full_text: str, chars_per_page: int = 3000
    ) -> list[dict]:
        """Build approximate page boundaries from paragraphs.

        DOCX doesn't have explicit page breaks, so we approximate
        based on character count.

        Args:
            paragraphs: List of paragraph info dicts
            full_text: Complete extracted text
            chars_per_page: Approximate characters per page

        Returns:
            List of page boundary dicts with number, start_char, end_char
        """
        if not full_text:
            return []

        pages: list[dict] = []
        text_length = len(full_text)
        page_number = 1
        current_pos = 0

        while current_pos < text_length:
            page_start = current_pos
            page_end = min(current_pos + chars_per_page, text_length)

            pages.append(
                {
                    "number": page_number,
                    "start_char": page_start,
                    "end_char": page_end,
                }
            )

            page_number += 1
            current_pos = page_end

        # If no pages were created (empty doc), return single empty page
        if not pages:
            pages.append(
                {
                    "number": 1,
                    "start_char": 0,
                    "end_char": 0,
                }
            )

        return pages
